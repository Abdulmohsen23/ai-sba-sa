from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.urls import reverse
from django.utils.translation import gettext as _
from django.http import HttpResponseRedirect

from .models import ProgramIdea, IdeaResponse , IdeaNote
from .forms import LanguageSelectionForm, StartIdeationForm, InitialConceptForm, ProgramDetailsForm , IdeaNoteForm
from .services import ProgramIdeationService
from core.utils import log_user_activity

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
from django.core.exceptions import PermissionDenied

from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io

from django.http import HttpResponse
from django.template.loader import render_to_string
from django.shortcuts import get_object_or_404
import io
from datetime import datetime


@login_required
def home(request):
    """Home page for Program Ideation tool."""
    return render(request, 'program_ideation/home.html', {
        'title': _('Program Ideation')
    })


@login_required
def select_language(request):
    """Select language for the ideation process."""
    if request.method == 'POST':
        form = LanguageSelectionForm(request.POST)
        if form.is_valid():
            language = form.cleaned_data['language']
            
            # Create a new program idea with the selected language
            idea = ProgramIdea.objects.create(
                user=request.user,
                language=language,
                current_step='start'
            )
            
            # Log activity
            log_user_activity(
                user=request.user,
                tool_name='program_ideation',
                action='select_language',
                details={'language': language},
                request=request
            )
            
            return redirect('program_ideation:start', idea_id=idea.id)
    else:
        form = LanguageSelectionForm()
    
    return render(request, 'program_ideation/select_language.html', {
        'form': form,
        'title': _('Select Language / اختر اللغة')
    })


@login_required
def start_ideation(request, idea_id):
    """Start the ideation process."""
    idea = get_object_or_404(ProgramIdea, id=idea_id, user=request.user)
    
    # Update translations based on language
    if idea.language == 'ar':
        title = "بدء صناعة البرنامج"
    else:
        title = "Start Program Creation"
    
    if request.method == 'POST':
        form = StartIdeationForm(request.POST)
        if form.is_valid():
            has_specific_idea = form.cleaned_data['has_specific_idea']
            idea.has_specific_idea = has_specific_idea
            idea.save()
            
            # Log activity
            log_user_activity(
                user=request.user,
                tool_name='program_ideation',
                action='start_ideation',
                details={'has_specific_idea': has_specific_idea},
                request=request
            )
            
            if has_specific_idea:
                idea.current_step = 'specific_idea'
                idea.save()
                return redirect('program_ideation:specific_idea', idea_id=idea.id)
            else:
                idea.current_step = 'no_specific_idea'
                idea.save()
                return redirect('program_ideation:no_specific_idea', idea_id=idea.id)
    else:
        form = StartIdeationForm()
    
    return render(request, 'program_ideation/start_ideation.html', {
        'form': form,
        'idea': idea,
        'title': title
    })


@login_required
def specific_idea(request, idea_id):
    """Handle the specific idea path."""
    idea = get_object_or_404(ProgramIdea, id=idea_id, user=request.user)
    
    # Ensure correct path
    if not idea.has_specific_idea:
        return redirect('program_ideation:start', idea_id=idea.id)
    
    # Update translations based on language
    if idea.language == 'ar':
        title = "تفاصيل الفكرة"
    else:
        title = "Idea Details"
    
    # Get or create the form with initial data
    if request.method == 'POST':
        form = ProgramDetailsForm(request.POST, instance=idea)
        if form.is_valid():
            form.save()
            
            # Log activity
            log_user_activity(
                user=request.user,
                tool_name='program_ideation',
                action='update_program_details',
                details={'idea_id': idea.id},
                request=request
            )
            
            # Check if all required fields are filled
            if idea.is_complete():
                # Generate discussion questions
                service = ProgramIdeationService(language=idea.language)
                discussion_questions = service.generate_discussion_questions(idea)
                
                # Save the response
                IdeaResponse.objects.create(
                    idea=idea,
                    response_type='discussion_questions',
                    content=discussion_questions
                )
                
                # Generate program format
                program_format = service.generate_program_format(idea)
                IdeaResponse.objects.create(
                    idea=idea,
                    response_type='program_format',
                    content=program_format
                )
                
                # Generate program script
                program_script = service.generate_program_script(idea)
                IdeaResponse.objects.create(
                    idea=idea,
                    response_type='program_script',
                    content=program_script
                )
                
                # Generate visual materials
                visual_materials = service.generate_visual_proposals(idea)
                IdeaResponse.objects.create(
                    idea=idea,
                    response_type='visual_materials',
                    content=visual_materials
                )
                
                # Update idea status
                idea.status = 'completed'
                idea.current_step = 'complete'
                idea.save()
                
                return redirect('program_ideation:complete', idea_id=idea.id)
            else:
                # Generate proposals for missing data
                service = ProgramIdeationService(language=idea.language)
                missing_data_proposals = service.get_missing_data_proposals(idea)
                
                # Save the response
                IdeaResponse.objects.create(
                    idea=idea,
                    response_type='missing_data',
                    content=missing_data_proposals
                )
                
                return redirect('program_ideation:missing_data', idea_id=idea.id)
    else:
        form = ProgramDetailsForm(instance=idea)
    
    return render(request, 'program_ideation/specific_idea.html', {
        'form': form,
        'idea': idea,
        'title': title
    })


@login_required
def missing_data(request, idea_id):
    """Show missing data proposals and continue with form."""
    idea = get_object_or_404(ProgramIdea, id=idea_id, user=request.user)
    
    # Get the latest missing data response
    response = IdeaResponse.objects.filter(
        idea=idea,
        response_type='missing_data'
    ).order_by('-created_at').first()
    
    # Update translations based on language
    if idea.language == 'ar':
        title = "اقتراحات للبيانات الناقصة"
    else:
        title = "Missing Data Suggestions"
    
    if request.method == 'POST':
        form = ProgramDetailsForm(request.POST, instance=idea)
        if form.is_valid():
            form.save()
            
            # Log activity
            log_user_activity(
                user=request.user,
                tool_name='program_ideation',
                action='update_program_details',
                details={'idea_id': idea.id},
                request=request
            )
            
            # Redirect back to specific idea to check for completeness
            return redirect('program_ideation:specific_idea', idea_id=idea.id)
    else:
        form = ProgramDetailsForm(instance=idea)
    
    return render(request, 'program_ideation/missing_data.html', {
        'form': form,
        'idea': idea,
        'response': response,
        'title': title
    })


@login_required
def no_specific_idea(request, idea_id):
    """Handle the no specific idea path."""
    idea = get_object_or_404(ProgramIdea, id=idea_id, user=request.user)
    
    # Ensure correct path
    if idea.has_specific_idea:
        return redirect('program_ideation:start', idea_id=idea.id)
    
    # Update translations based on language
    if idea.language == 'ar':
        title = "تصور أولي"
    else:
        title = "Initial Concept"
    
    if request.method == 'POST':
        form = InitialConceptForm(request.POST)
        if form.is_valid():
            has_initial_concept = form.cleaned_data['has_initial_concept']
            initial_concept = form.cleaned_data['initial_concept']
            
            idea.has_initial_concept = has_initial_concept
            if has_initial_concept:
                idea.initial_concept = initial_concept
            idea.save()
            
            # Log activity
            log_user_activity(
                user=request.user,
                tool_name='program_ideation',
                action='provide_initial_concept',
                details={
                    'has_initial_concept': has_initial_concept,
                    'idea_id': idea.id
                },
                request=request
            )
            
            # Generate ideas using the LLM service
            service = ProgramIdeationService(language=idea.language)
            
            if has_initial_concept:
                # Process the initial concept
                content = service.process_initial_concept(initial_concept)
            else:
                # Generate random ideas
                content = service.get_idea_suggestions()
            
            # Create response
            response = IdeaResponse.objects.create(
                idea=idea,
                response_type='suggestions',
                content=content
            )
            
            idea.current_step = 'suggestions'
            idea.save()
            
            return redirect('program_ideation:suggestions', idea_id=idea.id)
    else:
        form = InitialConceptForm()
    
    return render(request, 'program_ideation/no_specific_idea.html', {
        'form': form,
        'idea': idea,
        'title': title
    })


@login_required
def suggestions(request, idea_id):
    """Show suggestions and allow user to select one."""
    idea = get_object_or_404(ProgramIdea, id=idea_id, user=request.user)
    
    # Get the latest suggestions response
    response = IdeaResponse.objects.filter(
        idea=idea,
        response_type='suggestions'
    ).order_by('-created_at').first()
    
    # Update translations based on language
    if idea.language == 'ar':
        title = "اقتراحات البرامج"
        continue_button = "المتابعة بفكرة مختارة"
        restart_button = "البدء من جديد بفكرة خاصة بي"
    else:
        title = "Program Suggestions"
        continue_button = "Continue with Selected Idea"
        restart_button = "Start Over with My Own Idea"
    
    # Create a form instance for the template, similar to missing_data view
    form = ProgramDetailsForm(instance=idea)
    
    if request.method == 'POST':
        if 'continue' in request.POST:
            # Process the form data if submitted
            form = ProgramDetailsForm(request.POST, instance=idea)
            if form.is_valid():
                form.save()
            
            # User wants to continue with a selected idea
            idea.has_specific_idea = True
            idea.current_step = 'specific_idea'
            idea.save()
            
            # Log activity
            log_user_activity(
                user=request.user,
                tool_name='program_ideation',
                action='continue_with_suggestion',
                details={'idea_id': idea.id},
                request=request
            )
            
            return redirect('program_ideation:specific_idea', idea_id=idea.id)
        
        elif 'restart' in request.POST:
            # User wants to start over with their own idea
            idea.has_specific_idea = True
            idea.has_initial_concept = None
            idea.initial_concept = None
            idea.current_step = 'specific_idea'
            idea.save()
            
            # Log activity
            log_user_activity(
                user=request.user,
                tool_name='program_ideation',
                action='restart_with_own_idea',
                details={'idea_id': idea.id},
                request=request
            )
            
            return redirect('program_ideation:specific_idea', idea_id=idea.id)
    
    return render(request, 'program_ideation/suggestions.html', {
        'idea': idea,
        'response': response,
        'title': title,
        'continue_button': continue_button,
        'restart_button': restart_button,
        'form': form  # Add the form to the context
    })


@login_required
def complete(request, idea_id):
    """Show completed idea with all generated content."""
    idea = get_object_or_404(ProgramIdea, id=idea_id, user=request.user)
    
    # Get all responses
    responses = {}
    for response_type in ['discussion_questions', 'program_format', 'program_script', 'visual_materials']:
        response = IdeaResponse.objects.filter(
            idea=idea,
            response_type=response_type
        ).order_by('-created_at').first()
        
        if response:
            responses[response_type] = response
    
    # Update translations based on language
    if idea.language == 'ar':
        title = "اكتملت فكرة البرنامج"
        sections = {
            'discussion_questions': 'أسئلة النقاش',
            'program_format': 'تنسيق البرنامج',
            'program_script': 'نص البرنامج',
            'visual_materials': 'المقترحات البصرية'
        }
    else:
        title = "Program Idea Completed"
        sections = {
            'discussion_questions': 'Discussion Questions',
            'program_format': 'Program Format',
            'program_script': 'Program Script',
            'visual_materials': 'Visual Proposals'
        }
    
    return render(request, 'program_ideation/complete.html', {
        'idea': idea,
        'responses': responses,
        'sections': sections,
        'title': title
    })


@login_required
def idea_list(request):
    """List all program ideas."""
    ideas = ProgramIdea.objects.filter(user=request.user).order_by('-created_at')
    
    return render(request, 'program_ideation/idea_list.html', {
        'ideas': ideas,
        'title': _('My Program Ideas')
    })

# Add these to program_ideation/views.py

#  REPLACE the existing add_note view:
@login_required
def add_note(request, idea_id):
    """Enhanced add note view with multi-field support."""
    idea = get_object_or_404(ProgramIdea, id=idea_id, user=request.user)
    
    # Prepare field information
    fields = {}
    field_mapping = [
        ('program_name', 'Program Name', 'اسم البرنامج'),
        ('general_idea', 'General Idea', 'الفكرة العامة'),
        ('target_audience', 'Target Audience', 'الجمهور المستهدف'),
        ('program_objectives', 'Program Objectives', 'أهداف البرنامج'),
        ('program_type', 'Program Type', 'نوع البرنامج'),
        ('program_duration', 'Program Duration', 'مدة البرنامج'),
        ('episode_count', 'Episode Count', 'عدد الحلقات'),
        ('filming_location', 'Filming Location', 'موقع التصوير'),
    ]
    
    for field_id, en_label, ar_label in field_mapping:
        content = getattr(idea, field_id, '')
        fields[field_id] = {
            'label': ar_label if idea.language == 'ar' else en_label,
            'content': content[:100] + '...' if len(str(content)) > 100 else content,
            'is_empty': not bool(content),
        }
    
    if request.method == 'POST':
        selected_fields = request.POST.getlist('fields')
        note_content = request.POST.get('note_content', '').strip()
        note_type = request.POST.get('note_type', 'enhancement')
        priority = request.POST.get('priority', 'medium')
        
        if selected_fields and note_content:
            # Create the note
            note = IdeaNote.objects.create(
                idea=idea,
                note_content=note_content,
                note_type=note_type,
                priority=priority,
                status='pending'
            )
            
            # Create NoteField entries for each selected field
            from .models import NoteField
            for field_name in selected_fields:
                current_content = getattr(idea, field_name, '')
                NoteField.objects.create(
                    note=note,
                    field_name=field_name,
                    current_content=current_content
                )
            
            # Process the note
            process_note_task(note.id)
            
            messages.success(
                request,
                'تم إضافة الملاحظة بنجاح' if idea.language == 'ar' else 'Note added successfully'
            )
            return redirect('program_ideation:note_detail', note_id=note.id)
        else:
            messages.error(
                request,
                'يرجى تحديد حقل واحد على الأقل وكتابة الملاحظة' if idea.language == 'ar' 
                else 'Please select at least one field and write a note'
            )
    
    # Set title based on language
    if idea.language == 'ar':
        title = "إضافة ملاحظة"
    else:
        title = "Add Note"
    
    return render(request, 'program_ideation/add_note.html', {
        'idea': idea,
        'fields': fields,
        'title': title
    })


@login_required
def note_list(request, idea_id):
    """Show list of notes for an idea."""
    idea = get_object_or_404(ProgramIdea, id=idea_id, user=request.user)
    notes = idea.notes.all().order_by('-created_at')
    
    # Set title based on language
    if idea.language == 'ar':
        title = "ملاحظات البرنامج"
    else:
        title = "Program Notes"
    
    return render(request, 'program_ideation/note_list.html', {
        'idea': idea,
        'notes': notes,
        'title': title
    })


# UPDATE the note_detail view:
@login_required
def note_detail(request, note_id):
    """Enhanced note detail view."""
    note = get_object_or_404(IdeaNote, id=note_id)
    idea = note.idea
    
    # Ensure user owns the idea
    if idea.user != request.user:
        raise PermissionDenied
    
    # Get field content for backward compatibility
    field_content = ''
    if note.field_name:
        field_content = getattr(idea, note.field_name, '')
    
    # Set title based on language
    if idea.language == 'ar':
        title = "تفاصيل الملاحظة"
    else:
        title = "Note Details"
    
    return render(request, 'program_ideation/note_detail.html', {
        'note': note,
        'idea': idea,
        'field_content': field_content,
        'title': title,
    })


# ADD this new view for AJAX suggestion application:
@login_required
def apply_note_suggestion(request, note_id):
    """Apply a suggestion via AJAX."""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'message': 'Invalid method'})
    
    note = get_object_or_404(IdeaNote, id=note_id)
    idea = note.idea
    
    # Ensure user owns the idea
    if idea.user != request.user:
        return JsonResponse({'success': False, 'message': 'Permission denied'})
    
    try:
        data = json.loads(request.body)
        field_id = data.get('field_id')
        suggestion_text = data.get('suggestion_text')
        
        if not field_id or not suggestion_text:
            return JsonResponse({
                'success': False,
                'message': 'Missing field_id or suggestion_text'
            })
        
        # Apply the suggestion to the field
        if hasattr(idea, field_id):
            setattr(idea, field_id, suggestion_text)
            idea.save()
            
            # Track the application
            from .models import AppliedSuggestion
            AppliedSuggestion.objects.create(
                note=note,
                field_name=field_id,
                suggestion_text=suggestion_text
            )
            
            # Check if all fields have been addressed
            from .models import NoteField
            note_fields = note.note_fields.all()
            applied_fields = AppliedSuggestion.objects.filter(note=note).values_list('field_name', flat=True)
            
            if note_fields.exists():
                all_fields = [f.field_name for f in note_fields]
                if set(all_fields).issubset(set(applied_fields)):
                    note.is_applied = True
                    note.save()
            elif note.field_name and field_id == note.field_name:
                note.is_applied = True
                note.save()
            
            # Log activity
            log_user_activity(
                user=request.user,
                tool_name='program_ideation',
                action='apply_note_suggestion',
                details={
                    'idea_id': idea.id,
                    'note_id': note.id,
                    'field_name': field_id
                },
                request=request
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Suggestion applied successfully'
            })
        else:
            return JsonResponse({
                'success': False,
                'message': f'Invalid field: {field_id}'
            })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'message': str(e)
        })


# UPDATE the process_note_task function:
def process_note_task(note_id):
    """Process a note with multi-field support."""
    def _process():
        try:
            from .models import IdeaNote, NoteField
            from .services import ProgramIdeationService
            
            # Get the note
            note = IdeaNote.objects.get(id=note_id)
            note.status = 'processing'
            note.save()
            
            # Initialize service
            service = ProgramIdeationService(language=note.idea.language)
            
            # Check if note has multiple fields
            note_fields = note.note_fields.all()
            
            if note_fields.exists():
                # Process multi-field note
                response = service.process_multi_field_note(note)
            else:
                # Process single field note (backward compatibility)
                response = service.process_idea_note(note)
            
            # Update note with response
            note.response_content = response
            note.status = 'completed'
            note.save()
            
            return {'success': True, 'note_id': note.id}
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Error processing note ID {note_id}: {str(e)}")
            
            try:
                note = IdeaNote.objects.get(id=note_id)
                note.status = 'failed'
                note.save()
            except:
                pass
            
            return {'success': False, 'error': str(e)}
    
    # Start processing in a separate thread
    import threading
    thread = threading.Thread(target=_process)
    thread.daemon = True
    thread.start()
    
    return {'success': True, 'message': f'Processing started for note ID: {note_id}'}


# views.py


# PDF Export using ReportLab
@login_required
def export_idea_pdf(request, idea_id):
    """Export program idea as PDF"""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY, TA_RIGHT, TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    idea = get_object_or_404(ProgramIdea, id=idea_id, user=request.user)
    
    # Create response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="program_{idea.id}_{datetime.now().strftime("%Y%m%d")}.pdf"'
    
    # Create PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f2937'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    # Heading style
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#4b5563'),
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Arabic support (if needed)
    if idea.language == 'ar':
        # You need to add an Arabic font file
        # pdfmetrics.registerFont(TTFont('Arabic', 'path/to/arabic-font.ttf'))
        # Then use it in styles: fontName='Arabic'
        pass
    
    # Add title
    elements.append(Paragraph(idea.program_name or "Program Idea", title_style))
    elements.append(Spacer(1, 12))
    
    # Add metadata table
    metadata = [
        ['Created Date:', idea.created_at.strftime('%B %d, %Y')],
        ['Program Type:', idea.program_type or 'N/A'],
        ['Duration:', idea.program_duration or 'N/A'],
        ['Episodes:', idea.episode_count or 'N/A'],
    ]
    
    metadata_table = Table(metadata, colWidths=[2*inch, 4*inch])
    metadata_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.grey),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))
    
    elements.append(metadata_table)
    elements.append(Spacer(1, 20))
    
    # Add program details
    details = [
        ('General Idea', idea.general_idea),
        ('Target Audience', idea.target_audience),
        ('Program Objectives', idea.program_objectives),
        ('Filming Location', idea.filming_location),
    ]
    
    for label, content in details:
        if content:
            elements.append(Paragraph(f"<b>{label}:</b>", heading_style))
            elements.append(Paragraph(content, styles['BodyText']))
            elements.append(Spacer(1, 12))
    
    # Add page break before responses
    elements.append(PageBreak())
    
    # Add responses
    responses = IdeaResponse.objects.filter(idea=idea)
    
    for response in responses:
        elements.append(Paragraph(response.get_response_type_display(), heading_style))
        
        # Split content into paragraphs to handle long text
        content_lines = response.content.split('\n')
        for line in content_lines[:50]:  # Limit lines to prevent huge PDFs
            if line.strip():
                elements.append(Paragraph(line, styles['BodyText']))
        
        elements.append(Spacer(1, 20))
        
        # Add page break between major sections
        if response.response_type in ['program_format', 'program_script']:
            elements.append(PageBreak())
    
    # Build PDF
    doc.build(elements)
    
    # Get the value of the BytesIO buffer and write it to the response
    pdf = buffer.getvalue()
    buffer.close()
    response.write(pdf)
    
    return response


# Word Export using python-docx
@login_required
def export_idea_word(request, idea_id):
    """Export program idea as Word document"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    
    idea = get_object_or_404(ProgramIdea, id=idea_id, user=request.user)
    
    # Create document
    document = Document()
    
    # Set document properties
    document.core_properties.title = idea.program_name or "Program Idea"
    document.core_properties.author = request.user.get_full_name() or request.user.username
    
    # Add title
    title = document.add_heading(idea.program_name or "Program Idea", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata paragraph
    metadata = document.add_paragraph()
    metadata.add_run(f"Created: {idea.created_at.strftime('%B %d, %Y')}\n").bold = True
    metadata.add_run(f"Type: {idea.program_type or 'N/A'} | ")
    metadata.add_run(f"Duration: {idea.program_duration or 'N/A'} | ")
    metadata.add_run(f"Episodes: {idea.episode_count or 'N/A'}")
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph()  # Empty line
    
    # Add program details
    details = [
        ('General Idea', idea.general_idea),
        ('Target Audience', idea.target_audience),
        ('Program Objectives', idea.program_objectives),
        ('Program Type', idea.program_type),
        ('Program Duration', idea.program_duration),
        ('Episode Count', idea.episode_count),
        ('Filming Location', idea.filming_location),
    ]
    
    for label, content in details:
        if content:
            # Add heading
            heading = document.add_heading(label, level=2)
            
            # Add content
            para = document.add_paragraph(content)
            para.style = 'Normal'
            
            # Add space
            document.add_paragraph()
    
    # Add page break
    document.add_page_break()
    
    # Add responses
    responses = IdeaResponse.objects.filter(idea=idea)
    
    response_titles = {
        'discussion_questions': 'Discussion Questions',
        'program_format': 'Program Format',
        'program_script': 'Program Script',
        'visual_materials': 'Visual Materials',
        'suggestions': 'Suggestions',
        'missing_data': 'Missing Data Proposals'
    }
    
    for response in responses:
        # Add section title
        document.add_heading(response_titles.get(response.response_type, response.get_response_type_display()), level=1)
        
        # Add content
        for paragraph in response.content.split('\n'):
            if paragraph.strip():
                document.add_paragraph(paragraph)
        
        # Add page break between major sections
        if response.response_type in ['program_format', 'program_script']:
            document.add_page_break()
    
    # Save to response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="program_{idea.id}_{datetime.now().strftime("%Y%m%d")}.docx"'
    
    document.save(response)
    
    return response