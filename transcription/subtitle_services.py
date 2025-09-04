import os
import time
import whisper
from docx import Document
from datetime import datetime
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import torch
import warnings
import tempfile
from django.conf import settings
from django.utils import timezone
import logging
import subprocess
import json

warnings.filterwarnings("ignore")
logger = logging.getLogger(__name__)

# Enforce offline mode settings
os.environ["HF_HUB_DISABLE_TELEMETRY"] = "1"
os.environ["TRANSFORMERS_NO_ADVISORY_WARNINGS"] = "true"
os.environ["TOKENIZERS_PARALLELISM"] = "false"


class EnhancedSubtitleService:
    """Enhanced subtitle service with translation and subtitle generation."""
    
    def __init__(self):
        self.model = None
        self.translator_model = None
        self.tokenizer = None
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        
        # Subtitle configuration
        self.subtitle_config = {
            "mode": "subtitle",
            "max_chars_per_line": 42,
            "max_lines": 2,
            "min_duration": 1.7,
            "max_duration": 7.0,
            "reading_speed": 15,
            "gap_threshold": 0.5,
            "min_words": 4,
            "max_words": 25,
            "force_merge": True,
        }
    
    def load_models(self):
        """Load Whisper and translation models."""
        if self.model is None:
            logger.info("Loading Whisper large-v2 model...")
            try:
                self.model = whisper.load_model("large-v2")
                logger.info("✓ Whisper loaded successfully")
            except Exception as e:
                # Try with base model if large-v2 fails
                logger.warning(f"Could not load large-v2: {e}")
                logger.info("Loading Whisper base model instead...")
                self.model = whisper.load_model("base")
                logger.info("✓ Whisper base model loaded")
        
        # Try to load translation model (optional)
        try:
            if self.translator_model is None:
                from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
                
                model_name = "facebook/nllb-200-1.3B"
                logger.info(f"Loading NLLB-200 translator...")
                
                self.tokenizer = AutoTokenizer.from_pretrained(model_name)
                self.translator_model = AutoModelForSeq2SeqLM.from_pretrained(
                    model_name,
                    torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                )
                self.translator_model = self.translator_model.to(self.device)
                self.translator_model.eval()
                self.tokenizer.src_lang = "arb_Arab"
                logger.info("✓ NLLB-200 loaded successfully")
        except Exception as e:
            logger.warning(f"Could not load translation model: {e}")
            logger.info("Translation will not be available")
    
    def transcribe_video(self, video_path, source_language="ar"):
        """Transcribe video with word timestamps."""
        logger.info(f"Starting transcription of video: {video_path}")
        
        try:
            # Load models if not loaded
            self.load_models()
            
            # Transcribe with word timestamps
            result = self.model.transcribe(
                video_path,
                language=source_language,
                word_timestamps=True,
                verbose=False,
                task="transcribe",
                beam_size=5,
                best_of=5,
                fp16=False
            )
            
            return result
        except Exception as e:
            logger.error(f"Error during transcription: {str(e)}")
            raise
    
    def create_subtitle_segments(self, segments):
        """Create subtitle-ready segments with proper timing."""
        config = self.subtitle_config
        
        # First pass - merge very short segments
        merged_segments = []
        temp_segment = None
        
        for segment in segments:
            if temp_segment is None:
                temp_segment = segment.copy()
            else:
                gap = segment['start'] - temp_segment['end']
                combined_text = temp_segment['text'] + " " + segment['text']
                combined_duration = segment['end'] - temp_segment['start']
                
                if (gap < 1.0 and combined_duration <= config['max_duration'] * 1.2 and
                        len(combined_text) <= config['max_chars_per_line'] * config['max_lines'] * 1.2):
                    temp_segment['end'] = segment['end']
                    temp_segment['text'] = combined_text
                else:
                    merged_segments.append(temp_segment)
                    temp_segment = segment.copy()
        
        if temp_segment:
            merged_segments.append(temp_segment)
        
        # Second pass - ensure minimum segment length
        final_segments = []
        i = 0
        while i < len(merged_segments):
            segment = merged_segments[i]
            duration = segment['end'] - segment['start']
            word_count = len(segment['text'].split())
            
            if duration < config['min_duration'] or word_count < config.get('min_words', 5):
                if i + 1 < len(merged_segments):
                    next_segment = merged_segments[i + 1]
                    combined_text = segment['text'] + " " + next_segment['text']
                    combined_duration = next_segment['end'] - segment['start']
                    
                    if (combined_duration <= config['max_duration'] * 1.2 and
                            len(combined_text) <= config['max_chars_per_line'] * config['max_lines'] * 1.2):
                        final_segments.append({
                            'start': segment['start'],
                            'end': next_segment['end'],
                            'text': combined_text
                        })
                        i += 2
                        continue
            
            final_segments.append(segment)
            i += 1
        
        return final_segments
    
    def translate_text(self, text, source_lang="ar", target_lang="en"):
        """Translate text using NLLB model or return placeholder."""
        if not self.translator_model:
            return f"[Translation to {target_lang} not available]"
        
        try:
            # Set source and target languages
            if source_lang == "ar":
                self.tokenizer.src_lang = "arb_Arab"
                forced_bos_token_id = self.tokenizer.convert_tokens_to_ids("eng_Latn")
            else:
                self.tokenizer.src_lang = "eng_Latn"
                forced_bos_token_id = self.tokenizer.convert_tokens_to_ids("arb_Arab")
            
            inputs = self.tokenizer(
                text,
                return_tensors="pt",
                max_length=512,
                truncation=True,
                padding=True
            ).to(self.device)
            
            with torch.no_grad():
                generated_tokens = self.translator_model.generate(
                    **inputs,
                    forced_bos_token_id=forced_bos_token_id,
                    max_length=512,
                    num_beams=6,
                    length_penalty=1.0,
                    early_stopping=True,
                    temperature=0.7,
                    do_sample=False,
                    repetition_penalty=1.2,
                    no_repeat_ngram_size=3
                )
            
            translation = self.tokenizer.decode(generated_tokens[0], skip_special_tokens=True)
            return translation.strip()
            
        except Exception as e:
            logger.error(f"Translation error: {str(e)}")
            return f"[Translation error: {str(e)[:50]}]"
    
    def format_time_srt(self, seconds):
        """Format time for SRT format."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
    
    def format_time_vtt(self, seconds):
        """Format time for WebVTT format."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
    
    def create_srt_content(self, segments, use_translated=False):
        """Create SRT subtitle content."""
        lines = []
        for i, segment in enumerate(segments, 1):
            lines.append(str(i))
            lines.append(f"{self.format_time_srt(segment['start'])} --> {self.format_time_srt(segment['end'])}")
            
            if use_translated and 'translated_text' in segment:
                lines.append(segment['translated_text'])
            else:
                lines.append(segment['original_text'])
            lines.append("")
        
        return "\n".join(lines)
    
    def create_vtt_content(self, segments, use_translated=False):
        """Create WebVTT subtitle content."""
        lines = ["WEBVTT", ""]
        
        for i, segment in enumerate(segments, 1):
            lines.append(str(i))
            lines.append(f"{self.format_time_vtt(segment['start'])} --> {self.format_time_vtt(segment['end'])}")
            
            if use_translated and 'translated_text' in segment:
                lines.append(segment['translated_text'])
            else:
                lines.append(segment['original_text'])
            lines.append("")
        
        return "\n".join(lines)
    
    def get_video_duration(self, video_path):
        """Get video duration using ffprobe or fallback method."""
        try:
            cmd = [
                'ffprobe',
                '-v', 'error',
                '-show_entries', 'format=duration',
                '-of', 'default=noprint_wrappers=1:nokey=1',
                video_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            return float(result.stdout.strip())
            
        except Exception as e:
            logger.warning(f"Could not get duration with ffprobe: {e}")
            # Fallback: estimate from file size (rough approximation)
            try:
                file_size = os.path.getsize(video_path)
                # Assume average bitrate of 1 Mbps
                estimated_duration = (file_size * 8) / (1000000)  # Convert to seconds
                return estimated_duration
            except:
                return None


def create_subtitle_document(segments, filename, language, is_rtl=False, use_translated=False):
    """Create a Word document with subtitles."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f"{filename} - {'Translated' if use_translated else 'Original'} Subtitles", 0)
    
    # Metadata
    doc.add_paragraph(f"Language: {language.upper()}")
    doc.add_paragraph(f"Total Subtitles: {len(segments)}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # Add subtitles
    for i, segment in enumerate(segments, 1):
        # Format time
        start_min = int(segment['start'] // 60)
        start_sec = segment['start'] % 60
        end_min = int(segment['end'] // 60)
        end_sec = segment['end'] % 60
        
        # Add subtitle header
        header = f"[{start_min:02d}:{start_sec:05.2f} - {end_min:02d}:{end_sec:05.2f}] Subtitle {i}"
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(header)
        header_run.bold = True
        
        # Add subtitle text
        text = segment['translated_text'] if use_translated else segment['original_text']
        text_para = doc.add_paragraph(text)
        
        # Set RTL if needed
        if is_rtl:
            for para in [header_para, text_para]:
                p_format = para._element.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                p_format.append(bidi)
        
        doc.add_paragraph()
    
    return doc