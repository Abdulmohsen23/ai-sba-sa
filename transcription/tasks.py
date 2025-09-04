import os
import logging
from celery import shared_task
from django.conf import settings
from django.utils import timezone
from datetime import timedelta

logger = logging.getLogger(__name__)


@shared_task
def process_subtitle_generation(project_id):
    """Process subtitle generation for a video."""
    from transcription.models import VideoFile, SubtitleProject, SubtitleSegment, SubtitleStyle
    from transcription.subtitle_services import EnhancedSubtitleService
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    try:
        # Get project
        project = SubtitleProject.objects.get(id=project_id)
        video = project.video
        
        # Update status
        video.status = 'processing'
        video.save()
        
        logger.info(f"Starting subtitle generation for project {project_id}")
        
        # Initialize service
        service = EnhancedSubtitleService()
        
        # Get video path
        video_path = video.file.path
        
        # Transcribe video
        logger.info(f"Transcribing video: {video_path}")
        transcription_result = service.transcribe_video(
            video_path,
            source_language=project.source_language
        )
        
        # Create subtitle segments
        logger.info("Creating subtitle segments...")
        subtitle_segments = service.create_subtitle_segments(
            transcription_result['segments']
        )
        
        # Process each segment
        segments_data = []
        for idx, segment in enumerate(subtitle_segments, 1):
            original_text = segment['text'].strip()
            
            # Translate if needed
            translated_text = None
            if project.subtitle_mode == 'translate':
                logger.info(f"Translating segment {idx}/{len(subtitle_segments)}")
                translated_text = service.translate_text(
                    original_text,
                    source_lang=project.source_language,
                    target_lang=project.target_language
                )
            
            # Create SubtitleSegment in database
            SubtitleSegment.objects.create(
                project=project,
                segment_number=idx,
                start_time=segment['start'],
                end_time=segment['end'],
                original_text=original_text,
                translated_text=translated_text
            )
            
            segments_data.append({
                'start': segment['start'],
                'end': segment['end'],
                'original_text': original_text,
                'translated_text': translated_text
            })
        
        # Generate subtitle files
        logger.info("Generating subtitle files...")
        
        # Generate SRT files
        srt_content_original = service.create_srt_content(segments_data, use_translated=False)
        srt_filename_original = f"{video.id}_original.srt"
        srt_path_original = os.path.join(settings.MEDIA_ROOT, 'subtitles', srt_filename_original)
        os.makedirs(os.path.dirname(srt_path_original), exist_ok=True)
        
        with open(srt_path_original, 'w', encoding='utf-8') as f:
            f.write(srt_content_original)
        
        project.srt_file_arabic = f"subtitles/{srt_filename_original}"
        
        if project.subtitle_mode == 'translate':
            srt_content_translated = service.create_srt_content(segments_data, use_translated=True)
            srt_filename_translated = f"{video.id}_translated.srt"
            srt_path_translated = os.path.join(settings.MEDIA_ROOT, 'subtitles', srt_filename_translated)
            
            with open(srt_path_translated, 'w', encoding='utf-8') as f:
                f.write(srt_content_translated)
            
            project.srt_file_english = f"subtitles/{srt_filename_translated}"
        
        # Generate VTT files
        vtt_content_original = service.create_vtt_content(segments_data, use_translated=False)
        vtt_filename_original = f"{video.id}_original.vtt"
        vtt_path_original = os.path.join(settings.MEDIA_ROOT, 'subtitles', vtt_filename_original)
        
        with open(vtt_path_original, 'w', encoding='utf-8') as f:
            f.write(vtt_content_original)
        
        project.vtt_file_arabic = f"subtitles/{vtt_filename_original}"
        
        if project.subtitle_mode == 'translate':
            vtt_content_translated = service.create_vtt_content(segments_data, use_translated=True)
            vtt_filename_translated = f"{video.id}_translated.vtt"
            vtt_path_translated = os.path.join(settings.MEDIA_ROOT, 'subtitles', vtt_filename_translated)
            
            with open(vtt_path_translated, 'w', encoding='utf-8') as f:
                f.write(vtt_content_translated)
            
            project.vtt_file_english = f"subtitles/{vtt_filename_translated}"
        
        # Generate Word documents
        logger.info("Generating Word documents...")
        
        # Arabic document
        doc_arabic = create_subtitle_document(
            segments_data, 
            video.original_filename,
            project.source_language,
            is_rtl=(project.source_language == 'ar')
        )
        doc_filename_arabic = f"{video.id}_transcription_arabic.docx"
        doc_path_arabic = os.path.join(settings.MEDIA_ROOT, 'subtitles', doc_filename_arabic)
        doc_arabic.save(doc_path_arabic)
        project.doc_file_arabic = f"subtitles/{doc_filename_arabic}"
        
        # English document (if translated)
        if project.subtitle_mode == 'translate':
            doc_english = create_subtitle_document(
                segments_data,
                video.original_filename,
                'en',
                is_rtl=False,
                use_translated=True
            )
            doc_filename_english = f"{video.id}_transcription_english.docx"
            doc_path_english = os.path.join(settings.MEDIA_ROOT, 'subtitles', doc_filename_english)
            doc_english.save(doc_path_english)
            project.doc_file_english = f"subtitles/{doc_filename_english}"
        
        # Update project processing time
        project.processing_time = (timezone.now() - project.created_at).total_seconds() / 60
        project.save()
        
        # Update video status
        video.status = 'completed'
        
        # Set deletion date if retention is 5 days
        if video.retention == '5days':
            video.delete_at = timezone.now() + timedelta(days=5)
        
        video.save()
        
        logger.info(f"Subtitle generation completed for project {project_id}")
        
        return {
            'success': True,
            'project_id': project.id,
            'segments_count': len(segments_data)
        }
        
    except Exception as e:
        logger.error(f"Error processing project {project_id}: {str(e)}")
        
        # Update video status to failed
        try:
            video = VideoFile.objects.get(id=project.video.id)
            video.status = 'failed'
            video.error_message = str(e)
            video.save()
        except:
            pass
        
        return {
            'success': False,
            'error': str(e)
        }


def create_subtitle_document(segments, filename, language, is_rtl=False, use_translated=False):
    """Create a Word document with subtitles."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f"{filename} - {'Translated' if use_translated else 'Original'} Subtitles", 0)
    
    # Metadata
    doc.add_paragraph(f"Language: {language.upper()}")
    doc.add_paragraph(f"Total Subtitles: {len(segments)}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # Add subtitles
    for i, segment in enumerate(segments, 1):
        # Format time
        start_min = int(segment['start'] // 60)
        start_sec = segment['start'] % 60
        end_min = int(segment['end'] // 60)
        end_sec = segment['end'] % 60
        
        # Add subtitle header
        header = f"[{start_min:02d}:{start_sec:05.2f} - {end_min:02d}:{end_sec:05.2f}] Subtitle {i}"
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(header)
        header_run.bold = True
        
        # Add subtitle text
        text = segment['translated_text'] if use_translated else segment['original_text']
        text_para = doc.add_paragraph(text)
        
        # Set RTL if needed
        if is_rtl:
            for para in [header_para, text_para]:
                p_format = para._element.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                p_format.append(bidi)
        
        doc.add_paragraph()
    
    return doc


@shared_task
def export_video_with_subtitles(project_id):
    """Export video with burned-in subtitles."""
    from transcription.models import SubtitleProject
    from transcription.subtitle_services import EnhancedSubtitleService
    
    try:
        project = SubtitleProject.objects.get(id=project_id)
        service = EnhancedSubtitleService()
        
        # Get paths
        video_path = project.video.file.path
        srt_path = project.srt_file_arabic.path if project.srt_file_arabic else project.srt_file_english.path
        
        # Output path
        output_filename = f"{project.video.id}_subtitled.mp4"
        output_path = os.path.join(settings.MEDIA_ROOT, 'processed_videos', output_filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Get style settings
        style_settings = None
        if hasattr(project, 'style'):
            style = project.style
            style_settings = {
                'font_family': style.font_family,
                'font_size': style.font_size,
                'font_color': style.font_color,
                'alignment': 2 if style.text_align == 'center' else (1 if style.text_align == 'left' else 3)
            }
        
        # Burn subtitles
        success = service.burn_subtitles_to_video(
            video_path,
            srt_path,
            output_path,
            style_settings
        )
        
        if success:
            project.processed_video = f"processed_videos/{output_filename}"
            project.save()
            
            return {'success': True, 'output_path': output_path}
        else:
            return {'success': False, 'error': 'Failed to burn subtitles'}
            
    except Exception as e:
        logger.error(f"Error exporting video: {str(e)}")
        return {'success': False, 'error': str(e)}


@shared_task
def cleanup_old_videos():
    """Clean up videos that have passed their retention period."""
    from transcription.models import VideoFile
    
    now = timezone.now()
    
    # Find videos to delete
    videos_to_delete = VideoFile.objects.filter(
        delete_at__lte=now,
        delete_at__isnull=False
    )
    
    for video in videos_to_delete:
        try:
            # Delete file from storage
            if video.file:
                if os.path.exists(video.file.path):
                    os.remove(video.file.path)
            
            # Delete database record
            video.delete()
            
            logger.info(f"Deleted video {video.id} after retention period")
            
        except Exception as e:
            logger.error(f"Error deleting video {video.id}: {str(e)}")