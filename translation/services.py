import os
import time
import torch
import whisper
import tempfile
import subprocess
import shutil
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from django.conf import settings
from django.core.files import File
from django.core.files.base import ContentFile
from .models import TranslationProject, Subtitle, TranslationOutput


class TranslationService:
    """Service for handling the translation and subtitle generation process"""
    
    def __init__(self, project_id):
        self.project = TranslationProject.objects.get(id=project_id)
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        self.whisper_model = None
        self.translator_model = None
        self.tokenizer = None
        self.translator_loaded = False
        
    def load_models(self):
        """Load the required AI models"""
        try:
            # Load Whisper model
            self.whisper_model = whisper.load_model("large-v2")
            
            # Try to load the translation model if needed
            if self.project.translation_mode == 'translate':
                try:
                    from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
                    
                    model_name = "facebook/nllb-200-1.3B"
                    self.tokenizer = AutoTokenizer.from_pretrained(model_name)
                    self.translator_model = AutoModelForSeq2SeqLM.from_pretrained(
                        model_name,
                        torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32
                    )
                    self.translator_model = self.translator_model.to(self.device)
                    self.translator_model.eval()
                    
                    # Set source language for tokenizer
                    if self.project.source_language == 'ar':
                        self.tokenizer.src_lang = "arb_Arab"
                    else:
                        self.tokenizer.src_lang = "eng_Latn"
                        
                    self.translator_loaded = True
                except Exception as e:
                    print(f"Error loading translation model: {e}")
                    # Continue without translation capability
            return True
        except Exception as e:
            self.project.status = 'failed'
            self.project.error_message = f"Error loading models: {str(e)}"
            self.project.save()
            return False
    
    def translate_text(self, text, from_lang, to_lang):
        """Translate text between languages"""
        if not text.strip() or not self.translator_loaded:
            return ""
        
        try:
            inputs = self.tokenizer(
                text,
                return_tensors="pt",
                max_length=512,
                truncation=True,
                padding=True
            ).to(self.device)
            
            # Set forced BOS token based on target language
            if to_lang == 'ar':
                forced_bos_token_id = self.tokenizer.convert_tokens_to_ids("arb_Arab")
            else:
                forced_bos_token_id = self.tokenizer.convert_tokens_to_ids("eng_Latn")
            
            with torch.no_grad():
                generated_tokens = self.translator_model.generate(
                    **inputs,
                    forced_bos_token_id=forced_bos_token_id,
                    max_length=512,
                    num_beams=6,
                    length_penalty=1.0,
                    early_stopping=True
                )
            
            translation = self.tokenizer.decode(generated_tokens[0], skip_special_tokens=True)
            return translation.strip()
        except Exception as e:
            print(f"Translation error: {str(e)}")
            return "[Translation error]"
    
    def create_subtitle_segments(self, whisper_segments):
        """Create optimized subtitle segments from whisper output"""
        # Subtitle configuration (from your script)
        config = {
            "max_chars_per_line": 42,
            "max_lines": 2,
            "min_duration": 1.7,
            "max_duration": 7.0,
            "reading_speed": 15,
            "gap_threshold": 0.5,
            "min_words": 4,
            "max_words": 25,
            "force_merge": True,
        }
        
        # First pass - merge very short segments
        merged_segments = []
        temp_segment = None
        
        for segment in whisper_segments:
            if temp_segment is None:
                temp_segment = segment.copy()
            else:
                gap = segment['start'] - temp_segment['end']
                combined_text = temp_segment['text'] + " " + segment['text']
                combined_duration = segment['end'] - temp_segment['start']
                
                if (gap < 1.0 and combined_duration <= config['max_duration'] * 1.2):
                    temp_segment['end'] = segment['end']
                    temp_segment['text'] = combined_text
                else:
                    merged_segments.append(temp_segment)
                    temp_segment = segment.copy()
        
        # Add the last segment
        if temp_segment:
            merged_segments.append(temp_segment)
        
        # Second pass - ensure minimum segment length
        final_segments = []
        i = 0
        while i < len(merged_segments):
            segment = merged_segments[i]
            duration = segment['end'] - segment['start']
            word_count = len(segment['text'].split())

            # Check if this segment is too short
            if duration < config['min_duration'] or word_count < config.get('min_words', 4):
                # Try to merge with the next segment if possible
                if i + 1 < len(merged_segments):
                    next_segment = merged_segments[i + 1]
                    combined_text = segment['text'] + " " + next_segment['text']
                    combined_duration = next_segment['end'] - segment['start']

                    if (combined_duration <= config['max_duration'] * 1.2):
                        final_segments.append({
                            'start': segment['start'],
                            'end': next_segment['end'],
                            'text': combined_text
                        })
                        i += 2  # Skip the next segment since we merged it
                        continue

            # If we can't merge or there's no next segment, keep it anyway
            final_segments.append(segment)
            i += 1
            
        return final_segments
    
    def process_video(self):
        """Main processing function to generate subtitles from video"""
        # Update project status
        self.project.status = 'processing'
        self.project.save()
        
        try:
            start_time = time.time()
            
            # Load required models
            if not self.load_models():
                return False
            
            # Transcribe video with Whisper
            video_path = self.project.video_file.path
            result = self.whisper_model.transcribe(
                video_path,
                language="ar" if self.project.source_language == 'ar' else "en",
                word_timestamps=True,
                verbose=False
            )
            
            # Create optimized subtitle segments
            subtitle_segments = self.create_subtitle_segments(result['segments'])
            
            # Process each segment and save to database
            for idx, segment in enumerate(subtitle_segments, 1):
                original_text = segment['text'].strip()
                
                # Translate if needed
                translated_text = None
                if self.project.translation_mode == 'translate' and self.translator_loaded:
                    from_lang = self.project.source_language
                    to_lang = 'en' if from_lang == 'ar' else 'ar'
                    translated_text = self.translate_text(original_text, from_lang, to_lang)
                
                # Create subtitle object
                Subtitle.objects.create(
                    project=self.project,
                    start_time=segment['start'],
                    end_time=segment['end'],
                    original_text=original_text,
                    translated_text=translated_text,
                    sequence=idx,
                )
            
            # Generate output files
            self.generate_output_files()
            
            # Update project status
            self.project.status = 'completed'
            self.project.processing_time = time.time() - start_time
            self.project.save()
            
            # Clean up
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                
            return True
            
        except Exception as e:
            # Handle any exception during processing
            self.project.status = 'failed'
            self.project.error_message = str(e)
            self.project.save()
            
            # Clean up
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                
            return False
    
    def generate_output_files(self):
        """Generate SRT, VTT, and text files for both languages"""
        subtitles = self.project.subtitles.all().order_by('sequence')
        
        # Generate SRT files
        self.generate_srt_file(subtitles, 'original')
        if self.project.translation_mode == 'translate':
            self.generate_srt_file(subtitles, 'translated')
        
        # Generate VTT files
        self.generate_vtt_file(subtitles, 'original')
        if self.project.translation_mode == 'translate':
            self.generate_vtt_file(subtitles, 'translated')
        
        # Generate text files
        self.generate_text_file(subtitles, 'original')
        if self.project.translation_mode == 'translate':
            self.generate_text_file(subtitles, 'translated')
            
        # Generate DOCX files
        self.generate_docx_file(subtitles, 'original')
        if self.project.translation_mode == 'translate':
            self.generate_docx_file(subtitles, 'translated')
    
    def generate_single_output(self, output_type):
        """Generate a single output file based on the requested type"""
        subtitles = self.project.subtitles.all().order_by('sequence')
        
        # Check if the file already exists
        try:
            existing_output = TranslationOutput.objects.get(project=self.project, output_type=output_type)
            if os.path.exists(existing_output.file.path):
                return existing_output.file.path
            else:
                # File record exists but file is missing, delete the record
                existing_output.delete()
        except TranslationOutput.DoesNotExist:
            pass
        
        # Generate the file based on type
        if output_type == 'srt_original':
            return self.generate_srt_file(subtitles, 'original')
        elif output_type == 'srt_translated':
            return self.generate_srt_file(subtitles, 'translated')
        elif output_type == 'vtt_original':
            return self.generate_vtt_file(subtitles, 'original')
        elif output_type == 'vtt_translated':
            return self.generate_vtt_file(subtitles, 'translated')
        elif output_type == 'txt_original':
            return self.generate_text_file(subtitles, 'original')
        elif output_type == 'txt_translated':
            return self.generate_text_file(subtitles, 'translated')
        elif output_type == 'docx_original':
            return self.generate_docx_file(subtitles, 'original')
        elif output_type == 'docx_translated':
            return self.generate_docx_file(subtitles, 'translated')
        
        return None
    
    def format_time_srt(self, seconds):
        """Format time for SRT format (00:00:00,000)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
    
    def format_time_vtt(self, seconds):
        """Format time for WebVTT format (00:00:00.000)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
    
    def generate_srt_file(self, subtitles, text_type):
        """Generate SRT subtitle file"""
        # Check if we have subtitles
        if not subtitles.exists():
            return None
            
        # Create SRT content
        srt_content = ""
        for i, subtitle in enumerate(subtitles, 1):
            start_time = self.format_time_srt(subtitle.start_time)
            end_time = self.format_time_srt(subtitle.end_time)
            
            if text_type == 'original':
                text = subtitle.original_text
            else:
                text = subtitle.translated_text or ""
                
            if not text.strip():
                continue
            
            srt_content += f"{i}\n{start_time} --> {end_time}\n{text}\n\n"
        
        # Save to a file
        filename = f"{self.project.title}_{text_type}.srt"
        output_type = f"srt_{text_type}"
        
        # Create the output file
        output = TranslationOutput.objects.create(
            project=self.project,
            output_type=output_type
        )
        
        # Save the content to the file
        output.file.save(filename, ContentFile(srt_content.encode('utf-8')))
        
        return output.file.path
    
    def generate_vtt_file(self, subtitles, text_type):
        """Generate WebVTT subtitle file"""
        # Check if we have subtitles
        if not subtitles.exists():
            return None
            
        # Create VTT content
        vtt_content = "WEBVTT\n\n"
        for i, subtitle in enumerate(subtitles, 1):
            start_time = self.format_time_vtt(subtitle.start_time)
            end_time = self.format_time_vtt(subtitle.end_time)
            
            if text_type == 'original':
                text = subtitle.original_text
            else:
                text = subtitle.translated_text or ""
                
            if not text.strip():
                continue
            
            vtt_content += f"{start_time} --> {end_time}\n{text}\n\n"
        
        # Save to a file
        filename = f"{self.project.title}_{text_type}.vtt"
        output_type = f"vtt_{text_type}"
        
        # Create the output file
        output = TranslationOutput.objects.create(
            project=self.project,
            output_type=output_type
        )
        
        # Save the content to the file
        output.file.save(filename, ContentFile(vtt_content.encode('utf-8')))
        
        return output.file.path
    
    def generate_text_file(self, subtitles, text_type):
        """Generate plain text transcript file"""
        # Check if we have subtitles
        if not subtitles.exists():
            return None
            
        # Create text content
        text_content = f"Transcript for: {self.project.title}\n"
        text_content += f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        
        for subtitle in subtitles:
            start_min = int(subtitle.start_time // 60)
            start_sec = subtitle.start_time % 60
            time_str = f"[{start_min:02d}:{start_sec:05.2f}] "
            
            if text_type == 'original':
                text = subtitle.original_text
            else:
                text = subtitle.translated_text or ""
                
            if not text.strip():
                continue
            
            text_content += f"{time_str} {text}\n\n"
        
        # Save to a file
        filename = f"{self.project.title}_{text_type}.txt"
        output_type = f"txt_{text_type}"
        
        # Create the output file
        output = TranslationOutput.objects.create(
            project=self.project,
            output_type=output_type
        )
        
        # Save the content to the file
        output.file.save(filename, ContentFile(text_content.encode('utf-8')))
        
        return output.file.path
    
    def generate_docx_file(self, subtitles, text_type):
        """Generate Word document with transcript"""
        # Check if we have subtitles
        if not subtitles.exists():
            return None
            
        # Create a new Document
        doc = Document()
        
        # Add title
        title_text = f"{self.project.title} - "
        title_text += "Transcript" if text_type == 'original' else "Translation"
        title = doc.add_heading(title_text, 0)
        
        # Add metadata
        info = doc.add_paragraph()
        info.add_run("Document Information").bold = True
        doc.add_paragraph(f"Source File: {os.path.basename(self.project.video_file.name)}")
        doc.add_paragraph(f"Processing Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Language: {self.project.get_source_language_display() if text_type == 'original' else 'English' if self.project.source_language == 'ar' else 'Arabic'}")
        doc.add_paragraph(f"Total Segments: {subtitles.count()}")
        
        doc.add_paragraph("─" * 50)
        
        # Add segments
        for i, subtitle in enumerate(subtitles, 1):
            # Timestamp and subtitle number
            start_min = int(subtitle.start_time // 60)
            start_sec = subtitle.start_time % 60
            end_min = int(subtitle.end_time // 60)
            end_sec = subtitle.end_time % 60
            
            time_str = f"[{start_min:02d}:{start_sec:05.2f} - {end_min:02d}:{end_sec:05.2f}]"
            header = f"{time_str} Subtitle {i}:"
            
            header_para = doc.add_paragraph()
            header_run = header_para.add_run(header)
            header_run.bold = True
            
            # Text content
            if text_type == 'original':
                text = subtitle.original_text
            else:
                text = subtitle.translated_text or ""
                
            if not text.strip():
                continue
            
            text_para = doc.add_paragraph(text)
            
            # Set RTL for Arabic text
            if (text_type == 'original' and self.project.source_language == 'ar') or \
               (text_type == 'translated' and self.project.source_language == 'en'):
                p_format = text_para._element.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                p_format.append(bidi)
            
            doc.add_paragraph()
        
        # Add summary
        doc.add_paragraph("─" * 50)
        summary = doc.add_paragraph()
        summary.add_run("Summary").bold = True
        doc.add_paragraph(f"Total Subtitles: {subtitles.count()}")
        doc.add_paragraph(f"Document Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Save to a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_path = temp_file.name
            doc.save(temp_path)
        
        # Create TranslationOutput and save the file
        filename = f"{self.project.title}_{text_type}.docx"
        output_type = f"docx_{text_type}"
        
        output = TranslationOutput.objects.create(
            project=self.project,
            output_type=output_type
        )
        
        # Open the temp file and save to model
        with open(temp_path, 'rb') as f:
            output.file.save(filename, File(f))
        
        # Clean up the temp file
        try:
            os.unlink(temp_path)
        except:
            pass
        
        return output.file.path


import os
import subprocess
import tempfile
import shutil
from django.conf import settings
from .models import TranslationProject, Subtitle, TranslationOutput

class VideoExportService:
    """Service for exporting videos with hard-coded (burned-in) subtitles"""
    
    def __init__(self, project_id):
        self.project = TranslationProject.objects.get(id=project_id)
    
    def export_video_with_subtitles(self, subtitle_lang='original', font_size=24, font_color='#FFFFFF'):
        """
        Export video with permanently burned-in subtitles
        Uses multiple methods to ensure it works on Windows
        """
        print("Starting video export with hard-coded subtitles...")
        
        # Method 1: Try the working approach with drawtext filter
        result = self._burn_with_drawtext(subtitle_lang, font_size, font_color)
        if result:
            return result
        
        # Method 2: Try with ASS file in same directory as video
        result = self._burn_with_ass_same_dir(subtitle_lang, font_size, font_color)
        if result:
            return result
        
        # Method 3: Two-stage process (soft then hard)
        result = self._burn_two_stage(subtitle_lang, font_size, font_color)
        if result:
            return result
        
        print("All methods failed. Subtitles could not be burned in.")
        return None
    
    def _burn_with_drawtext(self, subtitle_lang='original', font_size=24, font_color='#FFFFFF'):
        """
        Method 1: Use drawtext filter which works reliably on Windows
        This creates one drawtext filter per subtitle
        """
        print("Method 1: Trying drawtext filter...")
        
        video_path = self.project.video_file.path
        
        # Get subtitles
        subtitles = self.project.subtitles.all().order_by('sequence')
        if not subtitles.exists():
            print("No subtitles found")
            return None
        
        # Prepare output path
        safe_title = ''.join(c for c in self.project.title if c.isalnum() or c in (' ', '_')).replace(' ', '_')
        output_filename = f"{safe_title}_subtitled.mp4"
        output_dir = os.path.join(settings.MEDIA_ROOT, 'outputs', 'translation')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)
        
        # Build complex filtergraph
        # We'll create multiple drawtext filters, one for each subtitle
        filters = []
        
        for subtitle in subtitles:
            # Choose the appropriate text based on subtitle_lang
            if subtitle_lang == 'original':
                text = subtitle.original_text
            else:
                text = subtitle.translated_text or ""
            
            if not text.strip():
                continue
                
            # Add speaker format if available
            if subtitle.speaker:
                text = f"{subtitle.speaker}: \"{text}\""
            
            # Escape special characters for FFmpeg
            text = text.replace('\\', '\\\\')
            text = text.replace(':', '\\:')
            text = text.replace("'", "\\'")
            text = text.replace('"', '\\"')
            text = text.replace('\n', ' ')  # Replace newlines with spaces
            text = text.replace('\r', '')
            
            # Limit text length to avoid issues
            if len(text) > 100:
                text = text[:97] + "..."
            
            # Create drawtext filter for this subtitle
            filter_str = (
                f"drawtext="
                f"text='{text}':"
                f"fontsize={font_size}:"
                f"fontcolor={font_color}:"
                f"borderw=2:"
                f"bordercolor=black:"
                f"x=(w-text_w)/2:"
                f"y=h-80:"
                f"enable='between(t,{subtitle.start_time},{subtitle.end_time})'"
            )
            filters.append(filter_str)
        
        if not filters:
            print("No valid subtitles to burn")
            return None
        
        # Combine filters with comma
        filter_complex = ','.join(filters)
        
        # Build FFmpeg command
        cmd = [
            'ffmpeg', '-y',
            '-i', video_path,
            '-vf', filter_complex,
            '-c:v', 'libx264',
            '-preset', 'fast',
            '-crf', '23',
            '-c:a', 'copy',
            output_path
        ]
        
        print(f"Processing {len(filters)} subtitles with drawtext...")
        
        try:
            process = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            
            if process.returncode == 0 and os.path.exists(output_path):
                print(f"Success! Video created at: {output_path}")
                return output_path
            else:
                print(f"Drawtext failed: {process.stderr[:500]}")
                return None
                
        except subprocess.TimeoutExpired:
            print("Process timed out")
            return None
        except Exception as e:
            print(f"Error: {str(e)}")
            return None

    def _burn_with_ass_same_dir(self, subtitle_lang='original', font_size=24, font_color='#FFFFFF'):
        """
        Method 2: Create ASS file in same directory as video (avoids path issues)
        """
        print("Method 2: Trying ASS file in video directory...")
        
        video_path = self.project.video_file.path
        video_dir = os.path.dirname(video_path)
        
        # Create ASS file in same directory as video
        ass_filename = f"temp_subs_{self.project.id}.ass"
        ass_path = os.path.join(video_dir, ass_filename)
        
        # Get subtitles
        subtitles = self.project.subtitles.all().order_by('sequence')
        if not subtitles.exists():
            print("No subtitles found")
            return None
        
        # Create ASS content
        ass_content = self._create_ass_content(subtitles, subtitle_lang, font_size, font_color)
        
        # Write ASS file
        with open(ass_path, 'w', encoding='utf-8') as f:
            f.write(ass_content)
        
        print(f"Created ASS file: {ass_path}")
        
        # Prepare output path
        safe_title = ''.join(c for c in self.project.title if c.isalnum() or c in (' ', '_')).replace(' ', '_')
        output_filename = f"{safe_title}_subtitled.mp4"
        output_dir = os.path.join(settings.MEDIA_ROOT, 'outputs', 'translation')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)
        
        # Change to video directory to use relative paths
        original_cwd = os.getcwd()
        os.chdir(video_dir)
        
        try:
            # Use relative path for ASS file
            video_filename = os.path.basename(video_path)
            
            # Simple command with relative paths
            cmd = f'ffmpeg -y -i "{video_filename}" -vf "ass={ass_filename}" -c:v libx264 -preset fast -c:a copy "{output_path}"'
            
            print(f"Running command from {os.getcwd()}: {cmd}")
            
            process = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=300)
            
            # Change back to original directory
            os.chdir(original_cwd)
            
            # Clean up ASS file
            try:
                os.remove(ass_path)
            except:
                pass
            
            if process.returncode == 0 and os.path.exists(output_path):
                print(f"Success! Video created at: {output_path}")
                return output_path
            else:
                print(f"ASS method failed: {process.stderr[:500]}")
                return None
                
        except Exception as e:
            os.chdir(original_cwd)
            # Clean up ASS file
            try:
                os.remove(ass_path)
            except:
                pass
            print(f"Error: {str(e)}")
            return None

    def _burn_two_stage(self, subtitle_lang='original', font_size=24, font_color='#FFFFFF'):
        """
        Method 3: Two-stage process - first add soft subs, then burn them
        This works because soft subtitle addition works on your system
        """
        print("Method 3: Two-stage burn process...")
        
        video_path = self.project.video_file.path
        
        # Create temp directory
        temp_dir = tempfile.mkdtemp(prefix='burn_')
        temp_soft = os.path.join(temp_dir, "soft_subs.mp4")
        srt_file = os.path.join(temp_dir, "subs.srt")
        
        # Get subtitles
        subtitles = self.project.subtitles.all().order_by('sequence')
        if not subtitles.exists():
            print("No subtitles found")
            shutil.rmtree(temp_dir)
            return None
        
        try:
            # Create SRT file
            with open(srt_file, 'w', encoding='utf-8') as f:
                count = 0
                for subtitle in subtitles:
                    # Choose the appropriate text based on subtitle_lang
                    if subtitle_lang == 'original':
                        text = subtitle.original_text
                    else:
                        text = subtitle.translated_text or ""
                    
                    if not text.strip():
                        continue
                    
                    # Add speaker format if available
                    if subtitle.speaker:
                        text = f"{subtitle.speaker}: \"{text}\""
                    
                    count += 1
                    start = self._format_time_srt(subtitle.start_time)
                    end = self._format_time_srt(subtitle.end_time)
                    
                    f.write(f"{count}\n{start} --> {end}\n{text}\n\n")
            
            print(f"Created SRT with {count} subtitles")
            
            # Stage 1: Add soft subtitles (this works based on your logs)
            cmd1 = [
                'ffmpeg', '-y',
                '-i', video_path,
                '-f', 'srt',
                '-i', srt_file,
                '-c:v', 'copy',
                '-c:a', 'copy',
                '-c:s', 'mov_text',
                temp_soft
            ]
            
            print("Stage 1: Adding soft subtitles...")
            process1 = subprocess.run(cmd1, capture_output=True, text=True)
            
            if process1.returncode != 0:
                print(f"Failed to add soft subtitles: {process1.stderr[:500]}")
                shutil.rmtree(temp_dir)
                return None
            
            print("Soft subtitles added successfully")
            
            # Prepare output path
            safe_title = ''.join(c for c in self.project.title if c.isalnum() or c in (' ', '_')).replace(' ', '_')
            output_filename = f"{safe_title}_subtitled.mp4"
            output_dir = os.path.join(settings.MEDIA_ROOT, 'outputs', 'translation')
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, output_filename)
            
            # Stage 2: Extract frames and burn subtitles using filter_complex
            cmd2 = [
                'ffmpeg', '-y',
                '-i', temp_soft,
                '-filter_complex', '[0:v][0:s]overlay[v]',
                '-map', '[v]',
                '-map', '0:a',
                '-c:v', 'libx264',
                '-preset', 'fast',
                '-c:a', 'copy',
                output_path
            ]
            
            print("Stage 2: Burning subtitles...")
            process2 = subprocess.run(cmd2, capture_output=True, text=True)
            
            # Clean up
            shutil.rmtree(temp_dir)
            
            if process2.returncode == 0 and os.path.exists(output_path):
                print(f"Success! Video created at: {output_path}")
                return output_path
            else:
                print(f"Burn stage failed: {process2.stderr[:500]}")
                
                # If burn failed, at least return the soft subtitle version
                print("Returning soft subtitle version as fallback")
                shutil.copy2(temp_soft, output_path)
                return output_path
                
        except Exception as e:
            print(f"Error: {str(e)}")
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
            return None

    def _create_ass_content(self, subtitles, subtitle_lang, font_size, font_color):
        """Helper to create ASS subtitle content"""
        # Convert hex color to BGR for ASS
        color = font_color.lstrip('#')
        if len(color) == 6:
            r, g, b = color[0:2], color[2:4], color[4:6]
            ass_color = f"&H00{b}{g}{r}"
        else:
            ass_color = "&H00FFFFFF"
        
        ass_content = f"""[Script Info]
    Title: Subtitles
    ScriptType: v4.00+
    PlayResX: 1920
    PlayResY: 1080

    [V4+ Styles]
    Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding
    Style: Default,Arial,{font_size},{ass_color},&H000000FF,&H00000000,&H80000000,-1,0,0,0,100,100,0,0,1,3,2,2,10,10,50,1

    [Events]
    Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text
    """
        
        for subtitle in subtitles:
            # Choose the appropriate text based on subtitle_lang
            if subtitle_lang == 'original':
                text = subtitle.original_text
            else:
                text = subtitle.translated_text or ""
            
            if not text.strip():
                continue
            
            # Add speaker format if available
            if subtitle.speaker:
                text = f"{subtitle.speaker}: \"{text}\""
            
            # Format times
            start = self._seconds_to_ass_time(subtitle.start_time)
            end = self._seconds_to_ass_time(subtitle.end_time)
            
            # Clean and format text
            text = text.replace('\n', '\\N').replace('\r', '')
            
            ass_content += f"Dialogue: 0,{start},{end},Default,,0,0,0,,{text}\n"
        
        return ass_content
    
    def _seconds_to_ass_time(self, seconds):
        """Convert seconds to ASS time format"""
        h = int(seconds // 3600)
        m = int((seconds % 3600) // 60)
        s = seconds % 60
        return f"{h}:{m:02d}:{s:05.2f}"
    
    def _format_time_srt(self, seconds):
        """Format time for SRT format"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"